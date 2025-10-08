import os
import re
import json
import time
import zipfile
import nbformat
from nbformat.v4 import new_notebook, new_markdown_cell
import concurrent.futures
from tqdm import tqdm
from PyPDF2 import PdfReader
import argparse
import requests
import sys
from typing import Optional, Dict, List, Tuple
from requests.adapters import HTTPAdapter
try:
    from requests.adapters import Retry  # type: ignore
except Exception:
    from urllib3.util.retry import Retry  # fallback
import subprocess  # For PDF export
import base64
import pandas as pd
from docx import Document

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DEFAULT_BASE_FOLDER = r"./udemyDownloads"
DEFAULT_AUTH_FILE = os.path.join(SCRIPT_DIR, "Authentication.json")
DEFAULT_CSS_FILE = os.path.join(SCRIPT_DIR, "custom.css")  # Optional custom CSS

def safe_name(name: str) -> str:
    s = re.sub(r'[<>:"/\\|?*\n\r\t]', "_", str(name))
    s = s.strip()
    s = re.sub(r"_+", "_", s)
    s = s.rstrip(" .")
    reserved = {
        "CON","PRN","AUX","NUL",
        "COM1","COM2","COM3","COM4","COM5","COM6","COM7","COM8","COM9",
        "LPT1","LPT2","LPT3","LPT4","LPT5","LPT6","LPT7","LPT8","LPT9"
    }
    if s.upper() in reserved or s == "":
        s = f"_{s}_" if s else "Untitled"
    if len(s) > 120:
        s = s[:120]
    return s

def _long_path(path: str) -> str:
    if os.name == "nt":
        ap = os.path.abspath(path)
        if ap.startswith("\\\\?\\"):
            return ap
        if ap.startswith("\\\\"):
            return "\\\\?\\UNC\\" + ap[2:]
        return "\\\\?\\" + ap
    return path

def is_texty(filename: str) -> bool:
    text_extensions = (
        ".txt", ".csv", ".tsv", ".json", ".jsonl", ".xml", ".html", ".htm", ".md", ".yaml", ".yml", ".ini", ".cfg",
        ".env", ".log", ".tex", ".rst", ".rmd", ".ipynb", ".epub",
        ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h", ".hpp", ".cs", ".go", ".rb", ".php", ".swift",
        ".kt", ".sql", ".sh", ".bat", ".ps1", ".r", ".pl", ".scala", ".dart", ".vue", ".svelte", ".css", ".scss", ".less",
        ".toml", "makefile", "dockerfile", ".graphql", ".graphqls", ".handlebars", ".mustache", ".jinja", ".jsp", ".asp", ".aspx"
    )
    fname = filename.lower()
    return fname.endswith(text_extensions) or fname in ("makefile", "dockerfile")

def is_image(filename: str) -> bool:
    return filename.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp"))

def is_video(filename: str) -> bool:
    return filename.lower().endswith((".mp4", ".avi", ".mkv", ".mov", ".webm"))

def is_pdf(filename: str) -> bool:
    return filename.lower().endswith('.pdf')

def is_zip(filename: str) -> bool:
    return filename.lower().endswith('.zip')

def is_excel(filename: str) -> bool:
    return filename.lower().endswith(('.xlsx', '.xls', '.ods'))

def is_docx(filename: str) -> bool:
    return filename.lower().endswith('.docx')

def extract_zip(zip_path, extract_to):
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            zip_ref.extractall(extract_to)
        return [os.path.join(extract_to, f) for f in zip_ref.namelist()]
    except Exception:
        return []

def preview_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        if reader.pages:
            text = reader.pages[0].extract_text()
            return text if text else "[No extractable text found in first page]"
        else:
            return "[No pages found in PDF]"
    except Exception as e:
        return f"[Error reading PDF: {e}]"

def export_notebook_to_pdf(notebook_path: str, pdf_path: str, css_path: Optional[str] = None):
    try:
        cmd = [
            sys.executable, "-m", "nbconvert",
            "--to", "webpdf",
            "--output", pdf_path,
            notebook_path
        ]
        subprocess.run(cmd, check=True)
        print(f"[pdf] Exported PDF: {pdf_path}")
    except Exception as e:
        print(f"[pdf] Failed to export PDF for {notebook_path}: {e}")

def resolve_base_and_downloads(base_folder: str) -> Tuple[str, str]:
    bf = os.path.abspath(base_folder)
    if os.path.basename(bf).lower() == "downloads":
        return os.path.dirname(bf), bf
    child = os.path.join(bf, "downloads")
    if os.path.isdir(child):
        return bf, child
    return bf, child

def load_transcriptions_map(json_path: str) -> Dict[Tuple[str, str, str], str]:
    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        out: Dict[Tuple[str, str, str], str] = {}
        for entry in data:
            cname = safe_name(entry.get("course_name", ""))
            sname = safe_name(entry.get("chapter_title", ""))
            lname = safe_name(entry.get("lecture_title", ""))
            trans = entry.get("transcription", "")
            if cname and sname and lname and trans:
                out[(cname, sname, lname)] = trans
        print(f"[transcriptions] Loaded {len(out)} transcription entries from {json_path}")
        return out
    except FileNotFoundError:
        print(f"[transcriptions] File not found: {json_path} (continuing without transcriptions)")
        return {}
    except Exception as e:
        print(f"[transcriptions] Failed to load transcriptions: {e} (continuing without transcriptions)")
        return {}

class UdemyApi:
    def __init__(self, auth_file="Authentication.json", user_id_hint="256172910", sleep_between_calls=0.05):
        self.auth_file = auth_file
        self.user_id_hint = user_id_hint
        self.sleep = sleep_between_calls
        self._load_auth()
        self._init_headers()
        self.session = self._init_session()

    def _load_auth(self):
        with open(self.auth_file, "r", encoding="utf-8") as f:
            data = json.load(f)
        self.ACCESS_TOKEN = data.get("access_token")
        if not self.ACCESS_TOKEN:
            raise ValueError("Missing 'access_token' in Authentication.json")

    def _init_headers(self):
        self.auth_header = {"Authorization": f"Bearer {self.ACCESS_TOKEN}", "Accept": "application/json"}
        self.cookie_headers = {
            "accept": "application/json",
            "x-requested-with": "XMLHttpRequest",
            "x-udemy-cache-user": str(self.user_id_hint)
        }
        self.cookies = {"access_token": self.ACCESS_TOKEN}

    def _init_session(self):
        session = requests.Session()
        retries = Retry(total=5, backoff_factor=0.5, status_forcelist=[429, 500, 502, 503, 504])
        session.mount('https://', HTTPAdapter(max_retries=retries))
        return session

    def get_course_name(self, course_id):
        url = f"https://www.udemy.com/api-2.0/courses/{course_id}/?fields[course]=title"
        r = self.session.get(url, headers=self.cookie_headers, cookies=self.cookies, timeout=30)
        return r.json().get("title", f"Course {course_id}") if r.status_code == 200 else f"Course {course_id}"

    def fetch_curriculum_map(self, course_id):
        url = (
           f"https://www.udemy.com/api-2.0/courses/{course_id}/subscriber-curriculum-items/?"
           f"curriculum_types=chapter,lecture,practice,quiz,role-play&page_size=200"
           f"&fields[lecture]=title,object_index,is_published,sort_order,created,asset,supplementary_assets,is_free"
           f"&fields[quiz]=title,object_index,is_published,sort_order,type"
           f"&fields[practice]=title,object_index,is_published,sort_order"
           f"&fields[chapter]=title,object_index,is_published,sort_order&fields[asset]=title,filename,asset_type,status,time_estimation,is_external&caching_intent=True"
        )
        r = self.session.get(url, headers=self.cookie_headers, cookies=self.cookies, timeout=30)
        r.raise_for_status()
        section_map, lecture_map = {}, {}
        current_section_id, current_section_title, current_section_idx = None, None, 0
        for item in r.json().get("results", []):
            if item.get("_class") == "chapter":
                current_section_id = item.get("id")
                current_section_title = item.get("title")
                current_section_idx = item.get("object_index", 0)
                section_map[current_section_id] = {"title": current_section_title, "index": current_section_idx}
            elif item.get("_class") == "lecture":
                lecture_map[item["id"]] = {
                    "lecture_title": item.get("title"),
                    "section_id": current_section_id,
                    "section_title": current_section_title,
                    "section_index": current_section_idx,
                    "lecture_index": item.get("object_index", 0),
                    "time_estimation": item.get("time_estimation"),
                    "supplementary_assets": item.get("supplementary_assets", [])
                }
        return section_map, lecture_map

    def _resolve_asset_url(self, course_id, lecture_id, sup_asset_id):
        url = (
            f"https://www.udemy.com/api-2.0/users/me/subscribed-courses/{course_id}"
            f"/lectures/{lecture_id}/supplementary-assets/{sup_asset_id}/"
            f"?fields[asset]=download_urls,time_estimation"
        )
        try:
            rr = self.session.get(url, headers=self.cookie_headers, cookies=self.cookies, timeout=30)
            rr.raise_for_status()
            data = rr.json()
            if "download_urls" in data and "File" in data["download_urls"]:
                return data["download_urls"]["File"][0]["file"], data.get("time_estimation")
            elif "asset" in data and "download_urls" in data["asset"]:
                return data["asset"]["download_urls"]["File"][0]["file"], data["asset"].get("time_estimation")
        except Exception:
            pass
        return None, None

    def enumerate_supplementary_assets(self, course_id, course_name):
        _, lecture_map = self.fetch_curriculum_map(course_id)
        rows = []
        for lecture_id, lec in lecture_map.items():
            supp = lec.get("supplementary_assets") or []
            if not supp:
                rows.append({
                    "course_id": course_id,
                    "course_name": course_name,
                    "section_id": lec.get("section_id"),
                    "section_name": lec.get("section_title"),
                    "section_index": lec.get("section_index", 0),
                    "lecture_id": lecture_id,
                    "lecture_name": lec.get("lecture_title"),
                    "lecture_index": lec.get("lecture_index", 0),
                    "asset_id": None,
                    "asset_title": None,
                    "download_url": None,
                    "time_estimation": lec.get("time_estimation"),
                    "is_stub": True
                })
                continue
            for asset in supp:
                sup_id = asset.get("id")
                asset_title = asset.get("title") or f"asset_{sup_id}"
                url, time_est = self._resolve_asset_url(course_id, lecture_id, sup_id)
                time.sleep(self.sleep)
                rows.append({
                    "course_id": course_id,
                    "course_name": course_name,
                    "section_id": lec.get("section_id"),
                    "section_name": lec.get("section_title"),
                    "section_index": lec.get("section_index", 0),
                    "lecture_id": lecture_id,
                    "lecture_name": lec.get("lecture_title"),
                    "lecture_index": lec.get("lecture_index", 0),
                    "asset_id": sup_id,
                    "asset_title": asset_title,
                    "download_url": url,
                    "time_estimation": time_est or lec.get("time_estimation"),
                    "is_stub": False
                })
        rows.sort(key=lambda r: (r.get("section_index", 0), r.get("lecture_index", 0), safe_name((r.get("asset_title") or "") or "")))
        return rows

def _parse_idx_and_name(folder_name: str) -> Tuple[int, str]:
    try:
        prefix, rest = folder_name.split("_", 1)
        idx = int(prefix)
        return idx, rest
    except Exception:
        return 0, folder_name

def _find_matching_child_dir(parent: str, target_name: Optional[str]) -> Optional[str]:
    if not os.path.isdir(parent):
        return None
    target_name = target_name or ""
    safe_target = safe_name(target_name)
    children = [d for d in os.listdir(parent) if os.path.isdir(os.path.join(parent, d))]
    for c in children:
        if c == target_name or c == safe_target:
            return os.path.join(parent, c)
    for c in children:
        parts = c.split("_", 1)
        if len(parts) == 2 and safe_name(parts[1]).lower() == safe_target.lower():
            return os.path.join(parent, c)
    for c in children:
        if c.lower() == target_name.lower() or safe_name(c).lower() == safe_target.lower():
            return os.path.join(parent, c)
    return None

def _list_lecture_files(base_folder: str, course_name: str, section_name: Optional[str], lecture_name: Optional[str]) -> List[Tuple[str, str]]:
    _, downloads_dir = resolve_base_and_downloads(base_folder)
    course_dir = _find_matching_child_dir(downloads_dir, course_name)
    if not course_dir:
        return []
    section_dir = _find_matching_child_dir(course_dir, section_name)
    if not section_dir:
        return []
    lecture_dir = _find_matching_child_dir(section_dir, lecture_name)
    if not lecture_dir:
        return []
    files: List[Tuple[str, str]] = []
    for fname in sorted(os.listdir(lecture_dir)):
        if fname.lower().endswith(".ipynb"):
            continue
        fpath = os.path.join(lecture_dir, fname)
        if os.path.isfile(fpath):
            files.append((fname, fpath))
    return files

def merge_api_rows_with_local(base_folder: str, course_name: str, api_rows: List[dict]) -> List[dict]:
    groups: Dict[Tuple[int, Optional[str], int, Optional[str]], List[dict]] = {}
    for r in api_rows:
        key = (r.get("section_index", 0), r.get("section_name"), r.get("lecture_index", 0), r.get("lecture_name"))
        groups.setdefault(key, []).append(r)

    merged: List[dict] = []
    for (s_idx, s_name, l_idx, l_name), rows in groups.items():
        local_files = _list_lecture_files(base_folder, course_name, s_name, l_name)
        file_map = {safe_name(fn): (fn, fp) for fn, fp in local_files}
        consumed = set()

        for r in rows:
            if r.get("is_stub"):
                merged.append(r)
                continue
            title = r.get("asset_title") or ""
            cand_keys = {title, safe_name(title)}
            chosen = None
            for k in cand_keys:
                ksafe = safe_name(k)
                if ksafe in file_map and ksafe not in consumed:
                    chosen = ksafe
                    break
            if not chosen and len(file_map) == 1 and not consumed:
                chosen = next(iter(file_map.keys()))
            if chosen:
                _, path = file_map[chosen]
                r = dict(r)
                r["local_path"] = path
                consumed.add(chosen)
            merged.append(r)

        for k, (fname, fpath) in file_map.items():
            if k in consumed:
                continue
            merged.append({
                "course_id": None,
                "course_name": course_name,
                "section_id": None,
                "section_name": s_name,
                "section_index": s_idx,
                "lecture_id": None,
                "lecture_name": l_name,
                "lecture_index": l_idx,
                "asset_id": None,
                "asset_title": fname,
                "download_url": None,
                "time_estimation": rows[0].get("time_estimation") if rows else None,
                "is_stub": False,
                "local_path": fpath
            })

        if not rows and not local_files:
            merged.append({
                "course_id": None,
                "course_name": course_name,
                "section_id": None,
                "section_name": s_name,
                "section_index": s_idx,
                "lecture_id": None,
                "lecture_name": l_name,
                "lecture_index": l_idx,
                "asset_id": None,
                "asset_title": None,
                "download_url": None,
                "time_estimation": None,
                "is_stub": True,
                "local_path": None
            })

    return merged

def _download_asset(url: str, save_path: str) -> bool:
    try:
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        with open(save_path, "wb") as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        print(f"[download] Asset downloaded: {save_path}")
        return True
    except Exception as e:
        print(f"[download] Failed to download asset from {url}: {e}")
        return False

class UdemyCourseNotebookBuilder:
    def __init__(self, base_folder, max_workers=16, transcriptions_map=None, css_path: Optional[str] = None):
        normalized_base, _ = resolve_base_and_downloads(base_folder)
        self.base_folder = normalized_base
        self.notebooks_dir = os.path.join(self.base_folder, "notebooks")
        os.makedirs(_long_path(self.notebooks_dir), exist_ok=True)
        self.max_workers = max_workers
        self.transcriptions_map = transcriptions_map or {}
        self.css_path = css_path

    def _notebook_path_for(self, course, section_idx, section, lecture_idx, lecture):
        section_folder = f"{section_idx:02d}_{safe_name(section or 'No Section')}"
        lecture_file = f"{lecture_idx:02d}_{safe_name(lecture or 'Untitled')}.ipynb"
        return os.path.join(self.notebooks_dir, safe_name(course), section_folder, lecture_file)

    def _build_lecture_notebook(self, course_name, section_idx, section_name, lecture_idx, lecture_name, rows):
        time_est = rows[0].get("time_estimation") if rows else None
        lecture_title = lecture_name or "Untitled"

        title_html = f"<h2 style='color:#1565c0;font-family:sans-serif;'>{lecture_title}</h2>"
        cells = [new_markdown_cell(title_html)]

        if time_est:
            duration_html = f"<b style='color:#1a237e;'>Duration:</b> <span style='font-size:14px;'>{time_est} min</span>"
            cells.append(new_markdown_cell(duration_html))

        if not rows or all(r.get("is_stub") for r in rows):
            cells.append(new_markdown_cell("<span style='color:#333;'>No supplementary assets for this lecture.</span>"))
        else:
            for r in rows:
                if r.get("is_stub"):
                    continue
                filename = r.get("asset_title") or "asset"
                url = r.get("download_url")
                lp = r.get("local_path")

                # Add asset link cell
                if url:
                    asset_link_html = f"<b style='color:#1565c0;'>Asset:</b> <a href='{url}' style='color:#0d47a1;'>{filename}</a>"
                else:
                    err = r.get("download_error") or "Unavailable"
                    asset_link_html = f"<b style='color:#1565c0;'>Asset:</b> {filename} <span style='color:red;'>(error: {err})</span>"
                cells.append(new_markdown_cell(asset_link_html))

                # Handle missing assets: download if possible
                if not lp and url:
                    target_dir = os.path.join(self.base_folder, "downloads", safe_name(course_name), safe_name(section_name or ""), safe_name(lecture_name or ""))
                    os.makedirs(_long_path(target_dir), exist_ok=True)
                    lp = os.path.join(target_dir, filename)
                    if _download_asset(url, lp):
                        r["local_path"] = lp

                # Append content based on file type
                if lp and os.path.exists(lp):
                    try:
                        if is_zip(filename):
                            extract_dir = lp + "_extracted"
                            os.makedirs(_long_path(extract_dir), exist_ok=True)
                            extracted_files = extract_zip(lp, extract_dir)
                            for ef in extracted_files:
                                ef_name = os.path.basename(ef)
                                try:
                                    if is_texty(ef_name) and os.path.getsize(ef) <= 512 * 1024:
                                        with open(_long_path(ef), "r", encoding="utf-8", errors="replace") as f:
                                            content = f.read()
                                        preview_html = f"<b style='color:#1565c0;'>Preview: {ef_name}</b><pre style='background:#f5f5f5;color:#263238;'>{content[:2000]}</pre>"
                                        cells.append(new_markdown_cell(preview_html))
                                    elif is_image(ef_name):
                                        with open(ef, "rb") as img_file:
                                            encoded_string = base64.b64encode(img_file.read()).decode('utf-8')
                                        img_html = f"<img src='data:image/png;base64,{encoded_string}' alt='{ef_name}' style='max-width:100%;'>"
                                        cells.append(new_markdown_cell(img_html))
                                    elif is_video(ef_name):
                                        video_html = f"<video controls style='max-width:100%;'><source src='{ef}' type='video/mp4'>Your browser does not support the video tag.</video>"
                                        cells.append(new_markdown_cell(video_html))
                                    elif is_excel(ef_name):
                                        try:
                                            df = pd.read_excel(ef)
                                            preview_html = f"<b style='color:#1565c0;'>Preview: {ef_name}</b><pre style='background:#f5f5f5;color:#263238;'>{df.head(10).to_string()}</pre>"
                                            cells.append(new_markdown_cell(preview_html))
                                        except Exception as e:
                                            error_html = f"<span style='color:red;'>Failed to preview Excel: {e}</span>"
                                            cells.append(new_markdown_cell(error_html))
                                    elif ef_name.lower().endswith('.docx'):
                                        try:
                                            doc = Document(ef)
                                            text = "\n".join([p.text for p in doc.paragraphs])
                                            preview_html = f"<b style='color:#1565c0;'>Preview: {ef_name}</b><pre style='background:#f5f5f5;color:#263238;'>{text[:2000]}</pre>"
                                            cells.append(new_markdown_cell(preview_html))
                                        except Exception as e:
                                            error_html = f"<span style='color:red;'>Failed to preview DOCX: {e}</span>"
                                            cells.append(new_markdown_cell(error_html))
                                    else:
                                        skip_html = f"<span style='color:#888;'>Preview not available for {ef_name} (binary or too large)</span>"
                                        cells.append(new_markdown_cell(skip_html))
                                except Exception as e:
                                    error_html = f"<span style='color:red;'>Failed to preview {ef_name}: {e}</span>"
                                    cells.append(new_markdown_cell(error_html))
                        elif is_pdf(filename):
                            pdf_text = preview_pdf(lp)
                            pdf_html = f"<b style='color:#1565c0;'>Preview of {filename} (first page):</b><pre style='background:#f5f5f5;color:#263238;'>{(pdf_text or '')[:2000]}</pre>"
                            cells.append(new_markdown_cell(pdf_html))
                        elif filename.lower().endswith(('.csv', '.tsv')):
                            try:
                                df = pd.read_csv(lp, sep=None, engine='python')
                                preview_html = f"<b style='color:#1565c0;'>Preview: {filename}</b><pre style='background:#f5f5f5;color:#263238;'>{df.head(10).to_string()}</pre>"
                                cells.append(new_markdown_cell(preview_html))
                            except Exception as e:
                                error_html = f"<span style='color:red;'>Failed to preview CSV/TSV: {e}</span>"
                                cells.append(new_markdown_cell(error_html))
                        elif is_excel(filename):
                            try:
                                df = pd.read_excel(lp)
                                preview_html = f"<b style='color:#1565c0;'>Preview: {filename}</b><pre style='background:#f5f5f5;color:#263238;'>{df.head(10).to_string()}</pre>"
                                cells.append(new_markdown_cell(preview_html))
                            except Exception as e:
                                error_html = f"<span style='color:red;'>Failed to preview Excel: {e}</span>"
                                cells.append(new_markdown_cell(error_html))
                        elif is_docx(filename):
                            try:
                                doc = Document(lp)
                                text = "\n".join([p.text for p in doc.paragraphs])
                                preview_html = f"<b style='color:#1565c0;'>Preview: {filename}</b><pre style='background:#f5f5f5;color:#263238;'>{text[:2000]}</pre>"
                                cells.append(new_markdown_cell(preview_html))
                            except Exception as e:
                                error_html = f"<span style='color:red;'>Failed to preview DOCX: {e}</span>"
                                cells.append(new_markdown_cell(error_html))
                        elif is_texty(filename) and os.path.getsize(lp) <= 512 * 1024:
                            with open(_long_path(lp), "r", encoding="utf-8", errors="replace") as f:
                                content = f.read()
                            preview_html = f"<b style='color:#1565c0;'>Preview: {filename}</b><pre style='background:#f5f5f5;color:#263238;'>{content[:2000]}</pre>"
                            cells.append(new_markdown_cell(preview_html))
                        elif is_image(filename):
                            with open(lp, "rb") as img_file:
                                encoded_string = base64.b64encode(img_file.read()).decode('utf-8')
                            img_html = f"<img src='data:image/png;base64,{encoded_string}' alt='{filename}' style='max-width:100%;'>"
                            cells.append(new_markdown_cell(img_html))
                        elif is_video(filename):
                            video_html = f"<video controls style='max-width:100%;'><source src='{lp}' type='video/mp4'>Your browser does not support the video tag.</video>"
                            cells.append(new_markdown_cell(video_html))
                        else:
                            skip_html = f"<span style='color:#888;'>Preview not available for {filename} (binary or unsupported type)</span>"
                            cells.append(new_markdown_cell(skip_html))
                    except Exception as e:
                        error_html = f"<span style='color:red;'>Failed to process {filename}: {e}</span>"
                        cells.append(new_markdown_cell(error_html))

        key = (safe_name(course_name), safe_name(section_name or ""), safe_name(lecture_name or ""))
        if key in self.transcriptions_map:
            transcription = self.transcriptions_map[key]
            trans_html = f"<b style='color:#1565c0;'>Transcription:</b><pre style='background:#f5f5f5;color:#263238;white-space:pre-wrap;'>{transcription[:5000]}</pre>"
            cells.append(new_markdown_cell(trans_html))

        return new_notebook(cells=cells)

    def build_notebooks_for_course(self, course_name: str, results: list) -> int:
        groups: Dict[Tuple[int, Optional[str], int, Optional[str]], List[dict]] = {}
        for r in results:
            key = (
                r.get("section_index", 0),
                r.get("section_name"),
                r.get("lecture_index", 0),
                r.get("lecture_name"),
            )
            groups.setdefault(key, []).append(r)

        created = 0
        with concurrent.futures.ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_nb = {
                executor.submit(self._build_lecture_notebook, course_name, s_idx, s_name, l_idx, l_name, rows): (s_idx, s_name, l_idx, l_name)
                for (s_idx, s_name, l_idx, l_name), rows in groups.items()
            }
            for future in tqdm(concurrent.futures.as_completed(future_to_nb), total=len(future_to_nb), desc=f"Building Notebooks for {course_name}", position=0):
                s_idx, s_name, l_idx, l_name = future_to_nb[future]
                nb = future.result()
                nb_path = self._notebook_path_for(course_name, s_idx, s_name, l_idx, l_name)
                target_dir = os.path.dirname(nb_path)
                os.makedirs(_long_path(target_dir), exist_ok=True)
                with open(_long_path(nb_path), "w", encoding="utf-8") as f:
                    nbformat.write(nb, f)
                pdf_path = nb_path.replace(".ipynb", ".pdf")
                export_notebook_to_pdf(nb_path, pdf_path, self.css_path)
                created += 1
        return created

def main(course_ids=None):
    parser = argparse.ArgumentParser(description="Build Udemy lecture notebooks for a single course from assets (API plan and/or downloads scan) with transcriptions.")
    parser.add_argument("--base-folder", default=DEFAULT_BASE_FOLDER, help="Base folder (parent of 'downloads'; notebooks will be created beside it).")
    parser.add_argument("--auth-file", default=DEFAULT_AUTH_FILE, help="Path to Authentication.json with access_token.")
    parser.add_argument("--api-plan", action="store_true", help="Fetch course/lecture/assets from Udemy API and build notebooks.")
    parser.add_argument("--merge-api-with-downloads", action="store_true", help="Attach local files to API assets per lecture.")
    parser.add_argument("--max-workers", type=int, default=16, help="Parallelism for notebook creation.")
    parser.add_argument("--css-file", default=DEFAULT_CSS_FILE, help="Custom CSS file for PDF styling.")
    parser.add_argument("--course-ids", nargs="+", type=int, help="Course IDs to process.")

    implicit_defaults = len(sys.argv) == 1
    args = parser.parse_args()

    if implicit_defaults:
        args.api_plan = True
        args.merge_api_with_downloads = True
        print("[main] No arguments supplied. Using implicit defaults: --api-plan --merge-api-with-downloads")
        print(f"[main] base-folder: {args.base_folder}")
        print(f"[main] auth-file: {args.auth_file}")

    normalized_base, downloads_dir = resolve_base_and_downloads(args.base_folder)
    print(f"[main] Base folder: {normalized_base}")
    print(f"[main] Downloads dir: {downloads_dir}")

    transcriptions_path = os.path.join(r"./udemyDownloads/videoTranscriptions.json")
    transcriptions_map = load_transcriptions_map(transcriptions_path)

    builder = UdemyCourseNotebookBuilder(
        base_folder=normalized_base,
        max_workers=args.max_workers,
        transcriptions_map=transcriptions_map,
        css_path=args.css_file
    )
    course_rows_map: Dict[str, List[dict]] = {}

    target_ids: List[int] = args.course_ids if args.course_ids else [2942646]
    if args.api_plan:
        api = UdemyApi(auth_file=args.auth_file)
        id_to_name = {}

        for cid in target_ids:
            cname = api.get_course_name(cid)
            print(f"[api] Planning course: {cname} ({cid})")
            api_rows = api.enumerate_supplementary_assets(cid, cname)
            rows = api_rows
            if args.merge_api_with_downloads:
                rows = merge_api_rows_with_local(normalized_base, cname, api_rows)
            course_rows_map[cname] = rows

    if not course_rows_map:
        print("[main] Nothing to build. Exiting.")
        return

    total_courses = len(course_rows_map)
    print(f"[build] Building notebooks for {total_courses} course(s).")
    for cname, rows in course_rows_map.items():
        created = builder.build_notebooks_for_course(cname, rows)
        print(f"[build] {cname}: created {created} notebook(s).")

if __name__ == "__main__":
    main()