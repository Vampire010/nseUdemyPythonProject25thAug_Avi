import os
import json
import requests
from openpyxl import load_workbook, Workbook
from docx import Document


class UdemyCourseDownloader:
    def __init__(self, base_folder, input_excel="udemy_courses.xlsx", auth_file="Authentication.json"):
        self.base_folder = base_folder
        self.input_excel = self.base_folder +"/"+input_excel
        self.auth_file = auth_file

        # Ensure base folder exists
        os.makedirs(self.base_folder, exist_ok=True)

        # Load authentication data
        with open(self.auth_file, "r") as f:
            auth_data = json.load(f)

        self.ACCESS_TOKEN = auth_data.get("access_token")
        self.CLIENT_ID = auth_data.get("client_id")
        self.CSRF = auth_data.get("csrf")

        self.cookies = {
            "access_token": self.ACCESS_TOKEN,
            "client_id": self.CLIENT_ID,
            "csrf": self.CSRF
        }

        self.headers = {
            'Accept': 'application/json, text/plain, */*',
            'X-Requested-With': 'XMLHttpRequest',
            'Cookie': '; '.join([f"{key}={value}" for key, value in self.cookies.items()])
        }

        # Prepare output files
        self.document = Document()
        self.document.add_heading("Udemy Course Curriculums", level=1)

        self.out_wb = Workbook()
        self.out_ws = self.out_wb.active
        self.out_ws.title = "Lecture Details"
        self.out_ws.append(["Course ID", "Course Name", "Lecture ID", "Lecture Title", "Object Index"])

    def fetch_courses(self):
        wb = load_workbook(self.input_excel)
        ws = wb.active

        for row in range(2, ws.max_row + 1):
            course_id = ws[f"B{row}"].value
            course_name = ws[f"C{row}"].value
            if not course_id:
                continue

            print(f"📘 Fetching Course ID: {course_id} - {course_name}")
            self.document.add_page_break()
            self.document.add_heading(f"Course ID: {course_id} - {course_name}", level=2)

            self.fetch_course_data(course_id, course_name)

    def fetch_course_data(self, course_id, course_name):
        url = (
            f"https://www.udemy.com/api-2.0/courses/{course_id}/subscriber-curriculum-items/?"
            "curriculum_types=chapter,lecture,practice,quiz,role-play&page_size=200"
            "&fields[lecture]=id,title,object_index,is_published,sort_order,created,asset,"
            "supplementary_assets,is_free"
            "&fields[quiz]=title,object_index,is_published,sort_order,type"
            "&fields[practice]=title,object_index,is_published,sort_order"
            "&fields[chapter]=title,object_index,is_published,sort_order"
            "&fields[asset]=title,filename,asset_type,status,time_estimation,is_external"
            "&caching_intent=True"
        )

        response = requests.get(url, headers=self.headers)
        if response.status_code != 200:
            self.document.add_paragraph(f"❌ Failed to fetch course. Status: {response.status_code}")
            return

        data = response.json()
        section_count = 0
        lecture_count = 0

        for item in data.get('results', []):
            if item['_class'] == 'chapter':
                section_count += 1
                object_index = item.get("object_index", "N/A")
                section_title = f"Section {section_count:02d} - {item['title']} (Index: {object_index})"
                self.document.add_heading(section_title, level=3)

            elif item['_class'] == 'lecture':
                lecture_count += 1
                lecture_id = item.get("id")
                lecture_title = item.get("title", "Untitled")
                object_index = item.get("object_index", "N/A")

                time_seconds = item.get('asset', {}).get('time_estimation')
                time_minutes = round(time_seconds / 60) if time_seconds else "N/A"
                time_str = f"{time_minutes} min" if time_seconds else "N/A"

                # Write to Word
                self.document.add_paragraph(f"{lecture_count:02d} - {lecture_title} (Time: {time_str})")

                # Write to Excel
                self.out_ws.append([course_id, course_name, lecture_id, lecture_title, object_index])

                # Resources
                if item.get('supplementary_assets'):
                    for asset in item['supplementary_assets']:
                        resource_title = asset.get('title', 'Untitled Resource')
                        self.document.add_paragraph(f"Resource: {resource_title}")

    def save_files(self):
        word_file = os.path.join(self.base_folder, "all_udemy_courses_curriculum.docx")
        excel_file = os.path.join(self.base_folder, "udemy_lecture_details.xlsx")

        print("\n💾 Saving Word and Excel files...")
        self.document.save(word_file)
        self.out_wb.save(excel_file)

        print(f"✅ Word file saved: {word_file}")
        print(f"✅ Excel file saved: {excel_file}")


if __name__ == "__main__":
    BASE_FOLDER = r"./udemyDownloads"

    udemy_downloader = UdemyCourseDownloader(base_folder=BASE_FOLDER)
    udemy_downloader.fetch_courses()
    udemy_downloader.save_files()
